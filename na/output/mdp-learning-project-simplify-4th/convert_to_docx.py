from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

with open('output/mdp-learning-project-simplify-4th/speech-script.md', 'r') as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')

    # Skip empty lines
    if not line.strip():
        i += 1
        continue

    # Skip horizontal rules
    if line.strip() == '---':
        i += 1
        continue

    # Title (# heading)
    if line.startswith('# '):
        text = line[2:].strip()
        p = doc.add_heading(text, level=1)
        i += 1
        continue

    # H2 (## heading) — part headers
    if line.startswith('## '):
        text = line[3:].strip()
        p = doc.add_heading(text, level=2)
        i += 1
        continue

    # Bold lines like **[Slide X — Title]** or **Slides 1-10...**
    if line.strip().startswith('**') and line.strip().endswith('**'):
        text = line.strip()[2:-2]
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        i += 1
        continue

    # Italic stage directions *(pause briefly)*
    if line.strip().startswith('*(') and line.strip().endswith(')*'):
        text = line.strip()[2:-2]
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(11)
        run.font.color.rgb = None
        i += 1
        continue

    # Regular paragraph
    text = line.strip()
    # Clean up any remaining markdown bold/italic markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    p = doc.add_paragraph(text)
    p.style.font.size = Pt(11)
    i += 1

out = 'output/mdp-learning-project-simplify-4th/speech-script.docx'
doc.save(out)
print(f'Saved to {out}')
